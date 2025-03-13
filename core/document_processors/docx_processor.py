# core/document_processors/docx_processor.py
import os
from typing import Dict, Any, List
import logging
from datetime import datetime
import docx

from .base_processor import BaseDocumentProcessor

logger = logging.getLogger(__name__)


class DocxProcessor(BaseDocumentProcessor):
    """Word文档处理器，用于处理DOCX文件"""

    def __init__(self, config=None):
        """
        初始化Word处理器

        Args:
            config: 可选配置参数
        """
        super().__init__(config or {})
        self.extract_tables = config.get("extract_tables", True)
        self.extract_headers_footers = config.get("extract_headers_footers", True)
        self.extract_comments = config.get("extract_comments", True)

        logger.info(
            f"初始化Word处理器: extract_tables={self.extract_tables}, extract_headers_footers={self.extract_headers_footers}")

    def extract_text(self, file_path: str) -> str:
        """
        从Word文档中提取文本

        Args:
            file_path: Word文档文件路径

        Returns:
            提取的文本内容
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")

        try:
            doc = docx.Document(file_path)
            content_parts = []

            # 提取正文段落
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            if paragraphs:
                content_parts.append("## 正文内容\n\n" + "\n\n".join(paragraphs))

            # 提取表格内容
            if self.extract_tables and doc.tables:
                tables_text = []
                for i, table in enumerate(doc.tables):
                    table_rows = []
                    for row in table.rows:
                        row_text = [cell.text.strip() for cell in row.cells]
                        table_rows.append(" | ".join(row_text))

                    tables_text.append(f"### 表格 {i + 1}\n" + "\n".join(table_rows))

                if tables_text:
                    content_parts.append("## 表格内容\n\n" + "\n\n".join(tables_text))

            # 提取页眉页脚
            if self.extract_headers_footers:
                headers_footers = []

                # 提取页眉
                try:
                    for section in doc.sections:
                        if section.header:
                            header_text = "\n".join([p.text for p in section.header.paragraphs if p.text.strip()])
                            if header_text:
                                headers_footers.append(f"页眉: {header_text}")

                        if section.footer:
                            footer_text = "\n".join([p.text for p in section.footer.paragraphs if p.text.strip()])
                            if footer_text:
                                headers_footers.append(f"页脚: {footer_text}")
                except Exception as hf_err:
                    logger.warning(f"提取页眉页脚时出错: {str(hf_err)}")

                if headers_footers:
                    content_parts.append("## 页眉页脚\n\n" + "\n".join(headers_footers))

            # 提取注释
            if self.extract_comments:
                comments = []
                try:
                    # python-docx库不直接支持提取注释，这需要额外处理
                    # 实际项目中可能需要使用其他库或方法
                    pass
                except Exception as comment_err:
                    logger.warning(f"提取注释时出错: {str(comment_err)}")

                if comments:
                    content_parts.append("## 注释\n\n" + "\n".join(comments))

            # 合并所有内容
            return "\n\n".join(content_parts)

        except Exception as e:
            logger.error(f"提取Word文本时出错: {str(e)}")
            raise

    def extract_metadata(self, file_path: str) -> Dict[str, Any]:
        """
        提取Word文档的元数据

        Args:
            file_path: Word文档文件路径

        Returns:
            包含文档元数据的字典
        """
        try:
            doc = docx.Document(file_path)
            metadata = {}

            # 提取文档属性
            try:
                core_props = doc.core_properties

                metadata.update({
                    "author": core_props.author,
                    "created": core_props.created.isoformat() if core_props.created else None,
                    "last_modified_by": core_props.last_modified_by,
                    "modified": core_props.modified.isoformat() if core_props.modified else None,
                    "title": core_props.title,
                    "subject": core_props.subject,
                    "keywords": core_props.keywords,
                    "comments": core_props.comments,
                    "category": core_props.category,
                    "version": core_props.revision
                })
            except Exception as prop_err:
                logger.warning(f"提取文档属性时出错: {str(prop_err)}")

            # 提取文档统计信息
            metadata.update({
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables),
                "section_count": len(doc.sections),
                "page_count": self._estimate_page_count(doc),
                "word_count": self._estimate_word_count(doc),
            })

            # 添加基本文件信息
            metadata.update({
                "file_size": os.path.getsize(file_path),
                "file_extension": "docx",
                "file_name": os.path.basename(file_path),
                "last_modified": datetime.fromtimestamp(
                    os.path.getmtime(file_path)
                ).isoformat()
            })

            return {k: v for k, v in metadata.items() if v is not None}

        except Exception as e:
            logger.error(f"提取Word元数据时出错: {str(e)}")
            raise

    @staticmethod
    def _estimate_page_count(doc) -> int:
        """
        估计文档的页数（近似值）

        Args:
            doc: docx文档对象

        Returns:
            估计的页数
        """
        try:
            # 实际页数需要渲染文档才能准确获得
            # 这里使用一个简单的估算方法
            total_para = len(doc.paragraphs)
            total_tables = len(doc.tables)
            total_chars = sum(len(p.text) for p in doc.paragraphs)

            # 假设平均每页2000字符，每个表格算200字符
            estimated_pages = max(1, (total_chars + total_tables * 200) // 2000)
            return estimated_pages
        except Exception:
            return 0

    def _estimate_word_count(self, doc) -> int:
        """
        估计文档的字数

        Args:
            doc: docx文档对象

        Returns:
            估计的字数
        """
        try:
            # 统计段落中的字数
            para_text = " ".join([p.text for p in doc.paragraphs])
            words = para_text.split()
            return len(words)
        except Exception:
            return 0

    def extract_styles(self, file_path: str) -> List[Dict[str, Any]]:
        """
        提取文档中使用的样式信息

        Args:
            file_path: Word文档文件路径

        Returns:
            样式信息列表
        """
        try:
            doc = docx.Document(file_path)
            styles = []

            # 收集使用的段落样式
            para_styles = set()
            for para in doc.paragraphs:
                if para.style:
                    para_styles.add(para.style.name)

            for style_name in sorted(para_styles):
                styles.append({
                    "name": style_name,
                    "type": "paragraph"
                })

            return styles
        except Exception as e:
            logger.error(f"提取Word样式时出错: {str(e)}")
            return []